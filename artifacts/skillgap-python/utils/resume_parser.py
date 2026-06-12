"""
Resume parser: extracts raw text from PDF and DOCX uploads.
Also detects structured sections (education, experience, certifications, projects).
"""

from __future__ import annotations

import io
import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ─── Section header patterns ──────────────────────────────────────────────────

SECTION_PATTERNS = {
    "contact": r"(?i)(contact|email|phone|linkedin|github|location|address)",
    "summary": r"(?i)(summary|objective|profile|about\s*me|professional\s*summary)",
    "experience": r"(?i)(experience|employment|work\s*history|career|positions?\s*held)",
    "education": r"(?i)(education|academic|degree|university|college|school)",
    "skills": r"(?i)(skills|technical\s*skills|core\s*competencies|expertise|proficiencies)",
    "projects": r"(?i)(projects?|portfolio|personal\s*projects?|side\s*projects?)",
    "certifications": r"(?i)(certifications?|certificates?|credentials?|licenses?)",
    "awards": r"(?i)(awards?|honors?|achievements?|accomplishments?|recognitions?)",
    "publications": r"(?i)(publications?|papers?|research|articles?)",
    "languages": r"(?i)(languages?|spoken\s*languages?)",
    "interests": r"(?i)(interests?|hobbies|activities)",
    "references": r"(?i)(references?|recommendations?)",
}

EDUCATION_KEYWORDS = [
    "bachelor", "master", "phd", "doctorate", "b.s.", "b.sc", "m.s.", "m.sc",
    "b.a.", "m.a.", "mba", "b.e.", "m.e.", "b.tech", "m.tech", "associate",
    "diploma", "certificate", "high school", "university", "college", "institute",
    "gpa", "cgpa", "graduated", "major", "minor", "coursework",
]

CERTIFICATION_KEYWORDS = [
    "certified", "certification", "certificate", "aws", "gcp", "azure",
    "pmp", "cka", "ckad", "cissp", "comptia", "oracle", "salesforce",
    "scrum", "agile", "itil", "tensorflow", "databricks", "coursera",
    "udemy", "edx", "google", "microsoft",
]

QUANTIFICATION_PATTERNS = [
    r"\d+[\+\%x]",
    r"\$[\d,]+[kKmM]?",
    r"\d+[\,\.]?\d*\s*(million|billion|thousand|k|m|b)\b",
    r"(increased|decreased|reduced|improved|boosted|saved)\s+by\s+\d+",
    r"\d+\s*(users|customers|clients|employees|team\s*members?|projects?|systems?)",
]


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file's bytes. Tries pdfplumber first, then PyPDF2."""
    text = ""

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n".join(pages)
        if text.strip():
            return _clean_text(text)
    except Exception as e:
        logger.warning("pdfplumber failed: %s — trying PyPDF2", e)

    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
        text = "\n".join(pages)
        if text.strip():
            return _clean_text(text)
    except Exception as e:
        logger.warning("PyPDF2 failed: %s", e)

    return ""


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file's bytes using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return _clean_text("\n".join(paragraphs))
    except Exception as e:
        logger.error("DOCX parsing failed: %s", e)
        return ""


def parse_file(uploaded_file) -> tuple[str, str]:
    """
    Parse an uploaded Streamlit file object.
    Returns (text, error_message). error_message is empty string on success.
    """
    if uploaded_file is None:
        return "", "No file provided."

    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if not file_bytes:
        return "", "The uploaded file appears to be empty."

    if filename.endswith(".pdf"):
        text = parse_pdf(file_bytes)
        if not text.strip():
            return "", "Could not extract text from the PDF. The file may be scanned/image-based or password-protected."
        return text, ""

    if filename.endswith(".docx"):
        text = parse_docx(file_bytes)
        if not text.strip():
            return "", "Could not extract text from the DOCX file. Please ensure it is a valid Word document."
        return text, ""

    if filename.endswith(".txt"):
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return _clean_text(text), ""
        except Exception as e:
            return "", f"Could not read text file: {e}"

    return "", f"Unsupported file format: {Path(filename).suffix}. Please upload PDF, DOCX, or TXT."


def extract_sections(text: str) -> dict[str, str]:
    """
    Split resume text into named sections.
    Returns a dict mapping section name → section text.
    """
    lines = text.split("\n")
    sections: dict[str, list[str]] = {"header": []}
    current_section = "header"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            sections.setdefault(current_section, []).append("")
            continue

        matched_section = _detect_section(stripped)
        if matched_section and len(stripped) < 60:
            current_section = matched_section
            sections.setdefault(current_section, [])
        else:
            sections.setdefault(current_section, []).append(stripped)

    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


def detect_contact_info(text: str) -> dict[str, str]:
    """Extract email, phone, LinkedIn, GitHub from text."""
    contact: dict[str, str] = {}

    email_match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    if email_match:
        contact["email"] = email_match.group(0)

    phone_match = re.search(
        r"(\+?1?\s?)?(\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4})", text
    )
    if phone_match:
        contact["phone"] = phone_match.group(0).strip()

    linkedin_match = re.search(
        r"(?:linkedin\.com/in/|linkedin:\s*)([a-zA-Z0-9\-]+)", text, re.IGNORECASE
    )
    if linkedin_match:
        contact["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"

    github_match = re.search(
        r"(?:github\.com/|github:\s*)([a-zA-Z0-9\-]+)", text, re.IGNORECASE
    )
    if github_match:
        contact["github"] = f"github.com/{github_match.group(1)}"

    return contact


def detect_education(text: str) -> list[str]:
    """Return lines that look like education entries."""
    lines = text.split("\n")
    edu_lines = []
    for line in lines:
        line_lower = line.lower()
        if any(kw in line_lower for kw in EDUCATION_KEYWORDS):
            stripped = line.strip()
            if stripped and len(stripped) > 5:
                edu_lines.append(stripped)
    return edu_lines[:10]


def detect_certifications(text: str) -> list[str]:
    """Return lines that look like certification entries."""
    lines = text.split("\n")
    cert_lines = []
    for line in lines:
        line_lower = line.lower()
        if any(kw in line_lower for kw in CERTIFICATION_KEYWORDS):
            stripped = line.strip()
            if stripped and len(stripped) > 5:
                cert_lines.append(stripped)
    return cert_lines[:15]


def detect_quantified_achievements(text: str) -> list[str]:
    """Return sentences/lines that contain quantified results."""
    lines = text.split("\n")
    quantified = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if any(re.search(pat, stripped, re.IGNORECASE) for pat in QUANTIFICATION_PATTERNS):
            quantified.append(stripped)
    return quantified[:15]


def analyze_resume_structure(text: str) -> dict:
    """
    High-level structural analysis of a resume.
    Returns section completeness, ATS risk factors, etc.
    """
    sections = extract_sections(text)
    contact = detect_contact_info(text)
    education = detect_education(text)
    certifications = detect_certifications(text)
    quantified = detect_quantified_achievements(text)

    word_count = len(text.split())
    has_summary = "summary" in sections
    has_experience = "experience" in sections
    has_skills = "skills" in sections
    has_education = "education" in sections or bool(education)
    has_contact = bool(contact)
    has_quantification = bool(quantified)

    completeness_score = sum([
        has_contact * 20,
        has_summary * 15,
        has_experience * 25,
        has_skills * 20,
        has_education * 15,
        has_quantification * 5,
    ])

    ats_risks = []
    if not has_contact:
        ats_risks.append("Missing contact information")
    if not has_summary:
        ats_risks.append("No professional summary/objective")
    if not has_skills:
        ats_risks.append("No dedicated skills section")
    if not has_quantification:
        ats_risks.append("No quantified achievements (numbers/percentages)")
    if word_count < 200:
        ats_risks.append("Resume is too short (< 200 words)")
    if word_count > 1200:
        ats_risks.append("Resume may be too long (> 1200 words)")

    return {
        "sections_found": list(sections.keys()),
        "contact_info": contact,
        "education_detected": education,
        "certifications_detected": certifications,
        "quantified_achievements": quantified,
        "word_count": word_count,
        "completeness_score": completeness_score,
        "ats_risks": ats_risks,
        "has_summary": has_summary,
        "has_experience": has_experience,
        "has_skills": has_skills,
        "has_education": has_education,
        "has_contact": has_contact,
    }


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    return text.strip()


def _detect_section(line: str) -> str | None:
    for section, pattern in SECTION_PATTERNS.items():
        if re.match(pattern, line):
            return section
    return None
